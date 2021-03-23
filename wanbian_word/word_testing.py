from docx import Document
from docx.shared import Inches


def CreateWord(dataQuestion, dataAnswer):
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('please review the questions carefully')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_page_break()

    document.save('data/doc/demo.docx')

def Output():
    print("我已经成功输出")