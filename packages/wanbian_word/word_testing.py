from docx import Document
from docx.shared import Inches
import docx
import random
import wanbian_hash.hash as wHash


class wordTesting:
    def __init__(self):
        # initialize extra class.
        self.WHash = wHash.Hash()

    # CreateHashArr() array the file name according to time hash value.
    def CreateHashArr(self, projectName, projectTitle, dataQuestion, dataAnswer, count):
        while count > 0:
            count -= 1
            document = Document()
            document.add_heading(projectTitle, 0)
            dict = {}
            dictNew = {}
            # put data in dict type.
            for i in range(0, len(dataQuestion)):
                dict[dataQuestion[i]] = dataAnswer[i]
            # random shuffles dict.
            dictKeysList = list(dict.keys())
            random.shuffle(dictKeysList)
            for k in dictKeysList:
                dictNew[k] = dict.get(k)
            # add those dict to the docx.
            for k, v in dictNew.items():
                document.add_paragraph(
                    k+"\n" + v + "\n", style='List Number'
                )
            document.add_page_break()
            document.save("data/"+projectName+"/"+self.WHash.hash()+".docx")

    # CreateSerialArr() array the file name according to serial number value.
    def CreateSerialArr(self, projectName, projectTitle, dataQuestion, dataAnswer, count):
        while count > 0:
            count -= 1
            document = Document()
            document.add_heading(projectTitle, 0)
            for i in range(0, len(dataQuestion)):
                document.add_paragraph(
                    dataQuestion[i]+"\n"+dataAnswer[i] + "\n", style='List Number'
                )
            document.add_page_break()
            document.save("data/"+projectName+"/" +
                          projectTitle+str(count)+".docx")
